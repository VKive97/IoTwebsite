import os
from bs4 import BeautifulSoup
from docx import Document

def extract_website_content(start_dir):
    doc = Document()
    doc.add_heading('Anstel Website Content', 0)
    
    exclude_dirs = {'.git', '.vscode', '.agents', '.claude', 'css', 'js', 'images', 'node_modules'}
    
    html_files = []
    for root, dirs, files in os.walk(start_dir):
        dirs[:] = [d for d in dirs if d not in exclude_dirs]
        for file in files:
            if file.endswith('.html'):
                html_files.append(os.path.join(root, file))
                
    html_files.sort()

    for file_path in html_files:
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
                
            soup = BeautifulSoup(content, 'html.parser')
            
            title = soup.title.string if soup.title else os.path.basename(file_path)
            if not title:
                title = os.path.basename(file_path)
            title = title.strip()
            
            for el in soup(['script', 'style', 'nav', 'header', 'footer', 'noscript', 'svg']):
                el.decompose()
                
            doc.add_heading(f"{title} - {os.path.relpath(file_path, start_dir)}", level=1)
            
            text = soup.get_text(separator='\n', strip=True)
            lines = [line.strip() for line in text.split('\n') if line.strip()]
            
            if lines:
                for line in lines:
                    doc.add_paragraph(line)
            else:
                doc.add_paragraph("[No readable text found]")
            
            doc.add_page_break()
            
        except Exception as e:
            print(f"Error processing {file_path}: {e}")

    output_file = os.path.join(start_dir, 'All_Pages_Content.docx')
    doc.save(output_file)
    print(f"Saved content to {output_file}")

if __name__ == '__main__':
    extract_website_content('.')
